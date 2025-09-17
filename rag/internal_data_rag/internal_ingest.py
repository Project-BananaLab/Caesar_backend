# -*- coding: utf-8 -*-
# 문서 임베딩 및 ChromaDB 저장 서비스
# internal_ingest.py

import os
import sys
import time
import zipfile
from pathlib import Path
from typing import List, Tuple, Optional

import pdfplumber
import docx
import openpyxl
from dotenv import load_dotenv

import chromadb
from chromadb.config import Settings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from openai import OpenAI

# ─────────────────────────────────────────────────────────
# 환경 변수 로드
# ─────────────────────────────────────────────────────────
load_dotenv()

# Chroma/Collection
CHROMA_PATH = os.getenv("CHROMA_PATH", "./chroma_data")
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "inside_data")

# 청킹 파라미터 (필요시 .env로 조절)
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "1000"))       # 청크 크기
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "150"))  # 오버랩

# 엑셀 폭발 방지 옵션
XLSX_MAX_ROWS_PER_SHEET = int(os.getenv("XLSX_MAX_ROWS_PER_SHEET", "10000"))
XLSX_MAX_COLS_PER_SHEET = int(os.getenv("XLSX_MAX_COLS_PER_SHEET", "512"))     # 🔹 추가: 열 상한 캡
XLSX_SKIP_HIDDEN_SHEETS = os.getenv("XLSX_SKIP_HIDDEN_SHEETS", "true").lower() == "true"

# 임베딩 요청 배치 한도 (요청당 토큰 상한 300k 대비 여유)
EMBED_MAX_TOKENS_PER_REQUEST = int(os.getenv("EMBED_MAX_TOKENS_PER_REQUEST", "280000"))
EMBED_MAX_ITEMS_PER_REQUEST = int(os.getenv("EMBED_MAX_ITEMS_PER_REQUEST", "256"))

# tiktoken은 선택적
try:
    import tiktoken
    _TIKTOKEN_ENC = tiktoken.get_encoding("cl100k_base")
except Exception:
    _TIKTOKEN_ENC = None

# OpenAI
client = OpenAI()


# ─────────────────────────────────────────────────────────
# 유틸: 실제 Office Open XML 포맷 스니핑(.docx/.xlsx 구분)
# ─────────────────────────────────────────────────────────
def _detect_office_kind(path: Path) -> Optional[str]:
    """
    ZIP 기반 Office 문서의 실제 종류를 추정:
      - 'docx'  : word/document.xml 존재
      - 'xlsx'  : xl/workbook.xml 존재
      - None    : ZIP 아님 또는 Office OpenXML 아님
    """
    try:
        if not zipfile.is_zipfile(path):
            return None
        with zipfile.ZipFile(path) as z:
            names = set(z.namelist())
        if any(n.startswith("word/") for n in names):
            return "docx"
        if any(n.startswith("xl/") for n in names):
            return "xlsx"
        return None
    except Exception:
        return None


# ─────────────────────────────────────────────────────────
# 임베딩 배치 유틸
# ─────────────────────────────────────────────────────────
def _estimate_tokens(text: str) -> int:
    """임베딩 토큰 대략치. tiktoken 있으면 정확, 없으면 문자수/4 근사."""
    if _TIKTOKEN_ENC is not None:
        try:
            return len(_TIKTOKEN_ENC.encode(text))
        except Exception:
            pass
    return max(1, len(text) // 4)


def embed_texts_batched(texts: List[str]) -> List[List[float]]:
    """토큰/아이템 예산을 지켜가며 여러 번으로 나눠 임베딩."""
    if not texts:
        return []

    batches: List[List[str]] = []
    current: List[str] = []
    current_tokens = 0

    for t in texts:
        tk = _estimate_tokens(t)

        # 단일 청크가 예산을 넘더라도(거의 없지만) 단독 배치로 보냄
        if tk > EMBED_MAX_TOKENS_PER_REQUEST:
            if current:
                batches.append(current)
                current, current_tokens = [], 0
            batches.append([t])
            continue

        if current and (
            current_tokens + tk > EMBED_MAX_TOKENS_PER_REQUEST
            or len(current) >= EMBED_MAX_ITEMS_PER_REQUEST
        ):
            batches.append(current)
            current, current_tokens = [], 0

        current.append(t)
        current_tokens += tk

    if current:
        batches.append(current)

    all_embeddings: List[List[float]] = []
    for i, batch in enumerate(batches, 1):
        print(f"  🔎 임베딩 배치 {i}/{len(batches)} (items={len(batch)}) 요청 중...")
        resp = client.embeddings.create(
            model="text-embedding-3-small",
            input=batch
        )
        all_embeddings.extend([d.embedding for d in resp.data])

    return all_embeddings


# ─────────────────────────────────────────────────────────
# 서비스 클래스
# ─────────────────────────────────────────────────────────
class IngestService:
    """문서 임베딩 및 ChromaDB 저장을 담당하는 서비스 클래스"""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", " ", ""]
        )
        self.supported_extensions = {".pdf", ".docx", ".xlsx"}

    # ========================= 파일 파싱 =========================
    def read_pdf(self, path: Path) -> str:  # PDF 파일 파싱
        texts = []
        try:
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        texts.append(t)
        except Exception as e:
            raise ValueError(f"PDF 로드 실패: {type(e).__name__}: {e}")
        return "\n\n".join(texts)

    def read_docx(self, path: Path) -> str:  # DOCX 파일 파싱
        try:
            d = docx.Document(str(path))
        except Exception as e:
            raise ValueError(f"DOCX 로드 실패: {type(e).__name__}: {e}")
        acc: List[str] = []
        acc.extend([p.text for p in d.paragraphs if p.text and p.text.strip()])
        # 테이블 추출(간단)
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    acc.append(" | ".join(cells))
        return "\n".join(acc)

    def read_xlsx(self, path: Path) -> str:  # XLSX 파일 파싱 (폭주 방지 트리밍/캡 적용)
        try:
            wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        except Exception as e:
            # 암호화/손상/비정상 구조 등 명확한 메시지 전달
            raise ValueError(f"엑셀 로드 실패: {type(e).__name__}: {e}")

        if not wb.worksheets:
            raise ValueError("엑셀에 워크시트가 없습니다.")

        acc: List[str] = []
        for ws in wb.worksheets:
            # 숨김 시트 스킵 옵션
            try:
                if XLSX_SKIP_HIDDEN_SHEETS and getattr(ws, "sheet_state", "visible") != "visible":
                    continue
            except Exception:
                pass

            acc.append(f"\n### [Sheet] {ws.title}")
            rows = 0

            # 🔹 열 상한 캡을 openpyxl 레벨에서 바로 적용
            iter_kwargs = {"values_only": True}
            if XLSX_MAX_COLS_PER_SHEET and XLSX_MAX_COLS_PER_SHEET > 0:
                iter_kwargs["max_col"] = XLSX_MAX_COLS_PER_SHEET

            for row in ws.iter_rows(**iter_kwargs):
                if rows >= XLSX_MAX_ROWS_PER_SHEET:
                    acc.append(f"...(truncated at {XLSX_MAX_ROWS_PER_SHEET} rows)")
                    break

                # 🔹 행 우측의 빈 열 트리밍: 실제 값이 있는 마지막 열까지만 사용
                last = -1
                # (열 캡이 적용된 범위 내에서만 검사)
                for i, v in enumerate(row):
                    sv = (str(v).strip() if v is not None else "")
                    if sv != "":
                        last = i

                if last < 0:
                    continue  # 완전 빈 행은 스킵

                # 🔹 최종 사용할 열 폭 결정
                width = last + 1
                if XLSX_MAX_COLS_PER_SHEET and XLSX_MAX_COLS_PER_SHEET > 0:
                    width = min(width, XLSX_MAX_COLS_PER_SHEET)

                # 🔹 최종 문자열 구성
                row_vals = []
                for v in row[:width]:
                    row_vals.append("" if v is None else str(v).strip())

                acc.append(" | ".join(row_vals))
                rows += 1

        return "\n".join(acc)

    def load_text(self, file_path: str, verbose: bool = True) -> str:
        """확장자 + 실제 포맷 스니핑으로 적절한 파서 선택"""
        p = Path(file_path)
        ext = p.suffix.lower()

        if verbose:
            print(f"  📄 파일 파싱 중: {p.name} ({ext})")

        actual = _detect_office_kind(p)  # 실제 포맷 스니핑(ZIP 기반 Office 문서의 실제 종류를 추정)

        try:
            if ext == ".pdf":   # PDF 파일 파싱
                return self.read_pdf(p)

            if ext == ".docx" or (actual == "docx" and ext != ".xlsx"):  # DOCX 파일 파싱
                if verbose and ext != ".docx" and actual == "docx":
                    print("  ⚠️ 확장자와 다른 실제 포맷(docx) 감지 → docx 파서 사용")
                return self.read_docx(p)

            if ext == ".xlsx" or (actual == "xlsx" and ext != ".docx"):  # XLSX 파일 파싱
                if verbose and ext != ".xlsx" and actual == "xlsx":
                    print("  ⚠️ 확장자와 다른 실제 포맷(xlsx) 감지 → xlsx 파서 사용")
                return self.read_xlsx(p)

            # 마지막 보루: 실제 포맷 기준 시도
            if actual == "docx":
                if verbose:
                    print("  ⚠️ 확장자 미지원/불명이나 실제 포맷(docx) 감지 → docx 파서 사용")
                return self.read_docx(p)
            if actual == "xlsx":
                if verbose:
                    print("  ⚠️ 확장자 미지원/불명이나 실제 포맷(xlsx) 감지 → xlsx 파서 사용")
                return self.read_xlsx(p)

            if verbose:
                print(f"  ⚠️ 지원하지 않는 파일 형식: {ext} (실제 포맷 미확인)")
            return ""

        except Exception as e:
            if verbose:
                print(f"  ❌ 파일 읽기 오류 ({p.name}): {e}")
            return ""

    # ========================= Chroma 헬퍼 =========================
    def get_chroma_collection(self):    # ChromaDB 컬렉션을 가져오거나 생성
        try:
            # ChromaDB 디렉토리 확인 및 생성
            Path(CHROMA_PATH).mkdir(parents=True, exist_ok=True)
            chroma = chromadb.PersistentClient(
                path=CHROMA_PATH,
                settings=Settings(
                    anonymized_telemetry=False,
                    is_persistent=True,
                ),
            )
            return chroma.get_or_create_collection(name=COLLECTION_NAME)
        except Exception as e:
            print(f"ChromaDB 초기화 오류: {str(e)}")
            print("새로운 ChromaDB 인스턴스로 재시도 중...")
            chroma = chromadb.Client()
            return chroma.get_or_create_collection(name=COLLECTION_NAME)

    # ========================= 단일 파일 처리 =========================
    def ingest_single_file(self, file_path: str, show_preview: bool = True) -> bool:
        print(f"📂 입력 파일: {file_path}")
        try:
            # 1) 파일 로드
            raw_text = self.load_text(file_path)
            if not raw_text.strip():
                print(f"❌ 빈 파일이거나 읽기 실패: {file_path}")
                return False

            print(f"✅ 파일 로드 완료, 전체 길이: {len(raw_text):,} chars")

            # 2) 텍스트 청킹
            chunks = self.text_splitter.split_text(raw_text)

            # 각 청크의 텍스트 길이 출력
            for i, c in enumerate(chunks):
                print(f"  [Chunk {i}] {len(c):,} chars")

            print(f"🪓 청킹 완료 → 총 {len(chunks)} chunks")
            if not chunks:
                print("❌ 청킹 결과가 비어 있습니다.")
                return False

            # 청크 미리보기
            if show_preview:
                for i, c in enumerate(chunks[:3]):
                    print(f"  [Chunk {i}] {c[:100]}...")

            # 3) 임베딩 생성
            print("⚙️ 임베딩 생성 중...")
            embeddings = embed_texts_batched(chunks)
            if not embeddings:
                print("❌ 임베딩 생성 실패(빈 입력).")
                return False
            print(f"✅ 임베딩 완료 → shape: {len(embeddings)} x {len(embeddings[0])}")

            # 4) ChromaDB 저장
            collection = self.get_chroma_collection()

            # 기존 동일 파일 청크 삭제(중복 방지)
            file_name = Path(file_path).name
            existing = collection.get(where={"source": file_name})
            if existing and existing.get("ids"):
                collection.delete(ids=existing["ids"])
                print(f"🗑 기존 {len(existing['ids'])} 청크 삭제")

            # 새 데이터 추가
            base_id = Path(file_path).stem
            ids = [f"{base_id}-{i}" for i in range(len(chunks))]
            metadatas = [{"source": file_name, "chunk_idx": i} for i in range(len(chunks))]

            collection.add(
                ids=ids,
                metadatas=metadatas,
                embeddings=embeddings,
                documents=chunks,
            )

            print(f"🎉 완료! {len(chunks)} chunks → Chroma collection '{COLLECTION_NAME}' 저장")
            return True

        except Exception as e:
            print(f"❌ 파일 처리 중 오류 발생: {str(e)}")
            return False

    # ========================= 다중 파일 처리 =========================
    def get_supported_files(self, folder_path: Path) -> List[Path]:  # 지원되는 파일 목록 추출
        files: List[Path] = []
        for file_path in folder_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                files.append(file_path)
        return sorted(files)    # 정렬된 파일 목록 반환(파일명 순)

    def process_single_file_batch(self, file_path: Path, collection) -> Tuple[int, bool]:  # 단일 파일 처리
        print(f"\n🔄 처리 중: {file_path.name}")                    #(처리된 청크 수, 성공 여부)
        try:
            # 1) 파일 로드
            raw_text = self.load_text(str(file_path))
            if not raw_text.strip():
                print(f"  ⚠️ 빈 파일이거나 읽기 실패: {file_path.name}")
                return 0, False

            print(f"  ✅ 파일 로드 완료, 전체 길이: {len(raw_text):,} chars")

            # 2) 텍스트 청킹
            chunks = self.text_splitter.split_text(raw_text)

            # 각 청크의 텍스트 길이 출력
            for i, c in enumerate(chunks):
                print(f"  [Chunk {i}] {len(c):,} chars")

            print(f"  🪓 청킹 완료 → 총 {len(chunks)} chunks")
            if not chunks:
                print(f"  ⚠️ 청킹 결과가 없음: {file_path.name}")
                return 0, False

            # 3) 임베딩 생성
            print(f"  ⚙️ 임베딩 생성 중... ({len(chunks)} chunks)")
            embeddings = embed_texts_batched(chunks)
            if not embeddings:
                print("  ⚠️ 임베딩 생성 실패(빈 입력)")
                return 0, False
            print(f"  ✅ 임베딩 완료 → shape: {len(embeddings)} x {len(embeddings[0])}")

            # 4) 기존 청크 삭제(중복 방지)
            file_name = file_path.name
            existing = collection.get(where={"source": file_name})
            if existing and existing.get("ids"):
                collection.delete(ids=existing["ids"])
                print(f"  🗑 기존 {len(existing['ids'])} 청크 삭제")

            # 5) 새 데이터 추가
            base_id = file_path.stem
            ids = [f"{base_id}-{i}" for i in range(len(chunks))]
            metadatas = [{"source": file_name, "chunk_idx": i} for i in range(len(chunks))]

            collection.add(
                ids=ids,
                metadatas=metadatas,
                embeddings=embeddings,
                documents=chunks,
            )

            print(f"  🎉 저장 완료! {len(chunks)} chunks → ChromaDB")
            return len(chunks), True

        except Exception as e:
            print(f"  ❌ 처리 오류 ({file_path.name}): {str(e)}")
            return 0, False

    # 폴더 내 모든 지원되는 파일들을 처리하여 ChromaDB에 저장
    def ingest_multiple_files(self, folder_path: str, clear_collection: bool = False) -> dict:
        folder = Path(folder_path)  # folder_path (str): 처리할 폴더 경로, clear_collection (bool): 처리 전 컬렉션 전체 삭제 여부 -> dict: 처리 결과 통계(성공/실패 파일 수, 총 청크 수, 소요 시간, 컬렉션 이름)
        if not folder.exists() or not folder.is_dir():
            print(f"❌ 폴더가 존재하지 않습니다: {folder_path}")
            return {"success": False, "error": "폴더가 존재하지 않음"}

        print(f"📂 폴더 처리 시작: {folder.absolute()}")

        files_to_process = self.get_supported_files(folder) # 지원되는 파일들 찾기
        if not files_to_process:
            print("❌ 처리할 수 있는 파일이 없습니다. (지원 형식: .pdf, .docx, .xlsx)")
            return {"success": False, "error": "처리할 파일이 없음"}

        print(f"📋 처리 대상 파일 {len(files_to_process)}개:")
        for i, file_path in enumerate(files_to_process, 1):
            file_size = file_path.stat().st_size
            size_mb = file_size / (1024 * 1024)
            print(f"  {i:2d}. {file_path.name} ({size_mb:.1f}MB)")

        print(f"\n🔧 ChromaDB 초기화 중... (경로: {CHROMA_PATH})")
        collection = self.get_chroma_collection()

        # 컬렉션 전체 삭제 옵션
        if clear_collection:
            try:
                existing_count = collection.count()
                if existing_count > 0:
                    print(f"🗑 기존 컬렉션 데이터 전체 삭제 중... ({existing_count} items)")
                    all_data = collection.get()
                    if all_data.get("ids"):
                        collection.delete(ids=all_data["ids"])
                    print("✅ 기존 데이터 삭제 완료")
            except Exception as e:
                print(f"⚠️ 기존 데이터 삭제 중 오류: {str(e)}")

        # 파일별 처리 통계
        total_chunks = 0
        successful_files = 0
        failed_files = 0
        start_time = time.time()

        print(f"\n🚀 파일 처리 시작... (총 {len(files_to_process)}개)")
        print("=" * 60)

        for i, file_path in enumerate(files_to_process, 1):
            print(f"\n[{i}/{len(files_to_process)}] {file_path.name}")

            chunks_count, success = self.process_single_file_batch(file_path, collection)

            if success:
                successful_files += 1
                total_chunks += chunks_count
            else:
                failed_files += 1

            progress = (i / len(files_to_process)) * 100    # 진행률 표시
            print(f"  📊 진행률: {progress:.1f}% ({i}/{len(files_to_process)})")

            if i < len(files_to_process):   # API 호출 제한을 위한 짧은 대기
                time.sleep(0.5)

        elapsed_time = time.time() - start_time

        result = {
            "success": True,
            "successful_files": successful_files,
            "failed_files": failed_files,
            "total_chunks": total_chunks,
            "elapsed_time": elapsed_time,
            "collection_name": COLLECTION_NAME,
        }

        print("\n" + "=" * 60)
        print("🎉 모든 파일 처리 완료!")
        print("📊 처리 결과:")
        print(f"  ✅ 성공: {successful_files}개 파일")
        print(f"  ❌ 실패: {failed_files}개 파일")
        print(f"  📝 총 청크 수: {total_chunks:,}개")
        print(f"  ⏱️ 소요 시간: {elapsed_time:.1f}초")
        print(f"  🗄️ 컬렉션: '{COLLECTION_NAME}'")
        print("=" * 60)

        return result


# ========================= 편의 함수 =========================
# 단일 파일 임베딩 편의 함수
def ingest_single_file(file_path: str, show_preview: bool = True) -> bool:
    return IngestService().ingest_single_file(file_path, show_preview)

# 다중 파일 임베딩 편의 함수
def ingest_multiple_files(folder_path: str, clear_collection: bool = False) -> dict:
    return IngestService().ingest_multiple_files(folder_path, clear_collection)


# ========================= CLI =========================
def main():
    print("=" * 80)
    print("📚 문서 임베딩 서비스")
    print("=" * 80)

    if len(sys.argv) < 2:
        print("사용법:")
        print("  단일 파일: python ingest_service.py <파일경로>")
        print("  다중 파일: python ingest_service.py <폴더경로> [--clear]")
        print("\n옵션:")
        print("  --clear: 처리 전 기존 컬렉션 데이터 전체 삭제")
        print("\n예시:")
        print("  python ingest_service.py ./rag/inside_data_rag/data/document.pdf")
        print("  python ingest_service.py ./rag/inside_data_rag/data")
        print("  python ingest_service.py ./rag/inside_data_rag/data --clear")
        sys.exit(1)

    path = sys.argv[1]
    clear_collection = "--clear" in sys.argv

    if clear_collection:
        print("⚠️ --clear 옵션이 설정되었습니다. 기존 데이터가 모두 삭제됩니다.")
        confirm = input("계속하시겠습니까? (y/N): ").strip().lower()
        if confirm not in ["y", "yes"]:
            print("❌ 작업이 취소되었습니다.")
            sys.exit(0)

    try:
        path_obj = Path(path)

        if path_obj.is_file():  # 단일 파일 처리
            print("📄 단일 파일 모드")
            success = ingest_single_file(path)
            sys.exit(0 if success else 1)

        elif path_obj.is_dir():  # 다중 파일 처리
            print("📂 다중 파일 모드")
            result = ingest_multiple_files(path, clear_collection)
            sys.exit(0 if result["success"] else 1)

        else:
            print(f"❌ 경로가 존재하지 않습니다: {path}")
            sys.exit(1)

    except KeyboardInterrupt:
        print("\n\n❌ 사용자에 의해 중단되었습니다.")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ 예상치 못한 오류 발생: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
